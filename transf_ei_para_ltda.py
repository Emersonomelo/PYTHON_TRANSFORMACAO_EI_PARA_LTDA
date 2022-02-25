from json import encoder
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import Inches
import os
import openpyxl
import pandas
import json
import sys
import urllib.request

pasta_nova = r"C:\\PYTHON_TRANSFORMACAO_EI_PARA_LTDA\\"
Document = Document()

def dados_da_receita(cnpj):
    
    url = 'http://receitaws.com.br/v1/cnpj/{0}'.format(cnpj)
    opener = urllib.request.build_opener()
    opener.addheaders = [
        ('User-agent',
        " Mozilla/5.0 (Windows NT 6.2; WOW64; rv:39.0) Gecko/20100101 Firefox/39.0")]

    with opener.open(url) as fd:
        content = fd.read().decode()

    dic = json.loads(content)
    nome = dic["nome"]
    planilha = openpyxl.Workbook()
    planilha.create_sheet('Planilha1teste', 0)
    planilha1 = planilha['Planilha1teste']

    planilha1['a1'].value ="abertura"
    planilha1['a2'].value ="situacao"
    planilha1['a3'].value ="tipo"
    planilha1['a4'].value ="nome"
    planilha1['a5'].value ="fantasia"
    planilha1['a6'].value ="porte"
    planilha1['a7'].value ="natureza_juridica"
    planilha1['a8'].value ="atividade_principal"
    planilha1['a9'].value ="atividades_secundarias"
    planilha1['a10'].value ="logradouro"
    planilha1['a11'].value ="numero"
    planilha1['a12'].value ="complemento"
    planilha1['a13'].value ="municipio"
    planilha1['a14'].value ="bairro"
    planilha1['a15'].value ="uf"
    planilha1['a16'].value ="cep"
    planilha1['a17'].value ="email"
    planilha1['a18'].value ="telefone"
    planilha1['a19'].value ="data_situacao"
    planilha1['a20'].value ="cnpj"
    planilha1['a21'].value ="ultima_atualizacao"
    planilha1['a22'].value ="status"
    planilha1['a23'].value ="efr"
    planilha1['a24'].value ="motivo_situacao"
    planilha1['a25'].value ="situacao_especial"
    planilha1['a26'].value ="data_situacao_especial"
    planilha1['a27'].value ="capital_social"
    planilha1['a28'].value ="qsa"
    planilha1['a29'].value ="extra"
    planilha1['a30'].value ="billing"
    planilha1['a31'].value ="cidade nascimento"
    planilha1['a32'].value ="data_de_nascimento" 
    planilha1['a33'].value ="estado_civil"
    planilha1['a34'].value ="profissão"
    planilha1['a35'].value ="rg"      
    planilha1['a36'].value ="expedido"
    planilha1['a37'].value ="CPF" 
    planilha1['a38'].value ="endereco"      
    planilha1['a39'].value ="mes" 
    planilha1['a40'].value ="ano" 
    planilha1['a41'].value ="data_desenquadramento"
    

    planilha1['b1'].value = str(dic["abertura"])
    planilha1['b2'].value = str(dic["situacao"])
    planilha1['b3'].value = str(dic["tipo"])
    planilha1['b4'].value = str(dic["nome"])
    planilha1['b5'].value = str(dic["fantasia"])
    planilha1['b6'].value = str(dic["porte"])
    planilha1['b7'].value = str(dic["natureza_juridica"])
    planilha1['b8'].value = str(dic["atividade_principal"][0]["text"])
    planilha1['b9'].value = str(dic["atividades_secundarias"])
    planilha1['b10'].value = str(dic["logradouro"])
    planilha1['b11'].value = str(dic["numero"])
    planilha1['b12'].value = str(dic["complemento"])
    planilha1['b13'].value = str(dic["municipio"])
    planilha1['b14'].value = str(dic["bairro"])
    planilha1['b15'].value = str(dic["uf"])
    planilha1['b16'].value = str(dic["cep"])
    planilha1['b17'].value = str(dic["email"])
    planilha1['b18'].value = str(dic["telefone"])
    planilha1['b19'].value = str(dic["data_situacao"])
    planilha1['b20'].value = str(dic["cnpj"])
    planilha1['b21'].value = str(dic["ultima_atualizacao"])
    planilha1['b22'].value = str(dic["status"])
    planilha1['b23'].value = str(dic["efr"])
    planilha1['b24'].value = str(dic["motivo_situacao"])
    planilha1['b25'].value = str(dic["situacao_especial"])
    planilha1['b26'].value = str(dic["data_situacao_especial"])
    planilha1['b27'].value = str(dic["capital_social"])
    planilha1['b28'].value = str(dic["qsa"])
    planilha1['b29'].value = str(dic["extra"])
    planilha1['b30'].value = str(dic["billing"])
    cidade_de_nascimento = input('Qual cidade de nascimento? ')
    planilha1['b31'].value = cidade_de_nascimento
    data_de_nascimento = input('Qual a data de nascimento? ')
    planilha1['b32'].value = data_de_nascimento
    estado_civil = input('Qual estado civil? ')
    planilha1['b33'].value = estado_civil
    planilha1['b34'].value ='Empresário'
    rg = input('Qual RG? ')
    planilha1['b35'].value = rg
    planilha1['b36'].value ='SSP/SP'
    cpf = str(dic["nome"])
    cpf = cpf[-11:]
    cpf = f'{cpf[0:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:11]}'
        
    planilha1['b37'].value = cpf
    rua = planilha1['b10'].value
    numero = planilha1['b11'].value
    complemento = planilha1['b12'].value
    municipio = planilha1['b13'].value
    bairro = planilha1['b14'].value
    uf = planilha1['b15'].value
    
    cep = str(dic["cep"])
    cep = f'{cep[0:3]}{cep[3:]}'
    
    planilha1['b38'].value = f'{rua}, {numero}, {complemento} Bairro: {bairro} - {municipio}, {uf} - CEP: {cep},'
    planilha1['b39'].value =input('Qual mes do desenquadramento? ') 
    planilha1['b40'].value =input('Qual ano do desenquadramento? ') 
    planilha1['b41'].value =str(input('Qual data de desenquadramento feito no site? '))
    
    rua = f'{rua}, {numero}, {complemento} Bairro: {bairro} - {municipio}, {uf} - CEP: {cep},'
    planilha.save(f'{pasta_nova}{nome}.xlsx')
    os.startfile(pasta_nova)
    return dic, rua

def transformacao_ei_em_ltda(cnpj, nome_cnpj, endereco_cnpj):
    #margens da pagina
    sections = Document.sections
    section = sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(2.5)
    section.bottom_margin = Inches(1)

    # Estilos:
    styles = Document.styles

    # Estilo do titulo
    titulo = styles.add_style("Titulo", WD_STYLE_TYPE.PARAGRAPH)
    titulo.font.name = "Arial"
    titulo.font.size = Pt(12)
    titulo.font.bold = False

    # Estilo do titulo com negrito
    titulo = styles.add_style("Titulo1", WD_STYLE_TYPE.PARAGRAPH)
    titulo.font.name = "Arial"
    titulo.font.size = Pt(12)
    titulo.font.bold = True

    # Estilo do paragrafo
    paragrafo = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    paragrafo.font.name = "Arial"
    paragrafo.font.size = Pt(12)
    paragrafo.font.bold = False

    # Estilo do paragrafo com italico
    paragrafo = styles.add_style("Paragraph4", WD_STYLE_TYPE.PARAGRAPH)
    paragrafo.font.name = "Arial"
    paragrafo.font.size = Pt(12)
    paragrafo.font.bold = False
    paragrafo.font.italic = True

    # Estilo do paragrafo com negrito
    paragrafo2 = styles.add_style("Paragraph2", WD_STYLE_TYPE.PARAGRAPH)
    paragrafo2.font.name = "Arial"
    paragrafo2.font.size = Pt(12)
    paragrafo2.font.bold = True

    T1 = Document.add_paragraph("CONSTITUIÇÃO POR TRANSFORMAÇÃO DE EMPRESÁRIO EM LTDA", style="Titulo")
    T1.alignment = 1

    nome_novo = 'CENTRO AUTOMOTIVO DULANS CAR LTDA' #input('Qual novo nome da LTDA? ').upper()
    p1 = Document.add_paragraph(f"{nome_novo}", style="Titulo1")
    p1.alignment = 1

    NIRE = '35850644622' #input('Qual numero do NIRE? ')

    p2 = Document.add_paragraph(f"NIRE: {NIRE}			CNPJ: {cnpj}", style="Paragraph")
    p2.alignment = 1

    nome_empresario = 'ALAN SANTOS RODRIGUES' #input('Qual nome do empresario? ')
    data_de_nascimento = '' #input('Qual é a data de nascimento? ')
    rg = '39.752.248-4 SSP-SP' #input('Qual RG e orgão emissor do empresario? ')
    cpf = '508.088.728-18' #input('Qual CPF do empresario? ')
    endereco = 'Avenida Fim de Semana, 998 - Jardim Casablanca – CEP: 05846-270 – São Paulo, SP' #input('Qual endereço do empresario? ')
    
       
    p3 = Document.add_paragraph(f'\n\t{nome_empresario}, Brasileiro (a), Solteiro(a), nascido (a) em {data_de_nascimento}, Empresário (a), portador (a) da cédula de identidade RG nº {rg} e do CPF {cpf}, domiciliado e residente no logradouro: {endereco}, responsável pela empresa estabelecida nesta praça sob a denominação social de {nome_cnpj} com sua sede na {endereco_cnpj} com requerimento de empresário arquivado na JUCESP sob o NIRE {NIRE} fazendo uso do que permite o § 3º do artigo 968 da lei 10.406/2002, com redação alterada pelo artigo 10 da lei complementar 128/2008, ora transforma seu registro de Empresário em SOCIEDADE EMPRESÁRIA LIMITADA passando a constituir o tipo jurídico Sociedade Empresaria limitada, a qual assume o ATIVO e PASSIVO da empresa individual e se regerá, doravante, pelo presente Contrato social o qual se obrigam mutuamente todos os sócios. ', style="Paragraph")
    p3.alignment = 3
    
    
    p4 = Document.add_paragraph(f'\nCLÁUSULA PRIMEIRA', style="Paragraph2")
    p4.alignment = 3
    
    p5 = Document.add_paragraph(f'\tA sociedade girará sob o nome empresarial {nome_novo} e sua sede {endereco_cnpj} registrada na receita federal CNPJ {cnpj}, a sociedade poderá a qualquer tempo abrir ou fechar filial, ou qualquer dependência, mediante alteração contratual deliberada na forma da lei. ', style="Paragraph")
    p5.alignment = 3
    
    p6 = Document.add_paragraph(f'\nCLÁUSULA SEGUNDA', style="Paragraph2")
    p6.alignment = 3
    
    p7 = Document.add_paragraph(f'\tO objeto será atividades de: ', style="Paragraph")
    p7.alignment = 3
    
    p8 = Document.add_paragraph(f'\nCLÁUSULA TERCEIRA', style="Paragraph2")
    p8.alignment = 3
    
    p9 = Document.add_paragraph(f'\tO início de suas atividades desde sua constituição como empresário em __/__/____ e seu prazo de duração é indeterminado.', style="Paragraph")
    p9.alignment = 3
    
    p10 = Document.add_paragraph(f'\nCLÁUSULA QUARTA', style="Paragraph2")
    p10.alignment = 3
        
    p11 = Document.add_paragraph(f'\tO capital social será de R$ 0.000,00 ( mil reais) dividido em 0.000,00 ( mil) quotas sociais no valor nominal de R$ 1,00 (Um real) cada uma, totalmente subscrito e integralizado em Moeda Corrente do país, para o (a) sócio (a) {nome_empresario}, acima qualificado (a). ', style="Paragraph")
    p11.alignment = 3
    
    p12 = Document.add_paragraph(f'\nCLÁUSULA QUINTA', style="Paragraph2")
    p12.alignment = 3
    
    p13 = Document.add_paragraph(f'\tAs quotas são indivisíveis e não poderão ser cedidas ou transferidas a terceiros sem o consentimento do(s) outro(s) sócio(s), a quem fica assegurado, em igualdade de condições e preço direito de preferência para a sua aquisição se postas à venda, formalizando, se realizada a cessão, a alteração contratual pertinente. ', style="Paragraph")
    p13.alignment = 3    
    
    p14 = Document.add_paragraph(f'\nCLÁUSULA SEXTA', style="Paragraph2")
    p14.alignment = 3
    
    p15 = Document.add_paragraph(f'\tA responsabilidade de cada sócio é restrita ao valor de suas quotas, mas todos respondem solidariamente pela integralização do capital social. ', style="Paragraph")
    p15.alignment = 3  
 
    p16 = Document.add_paragraph(f'\nCLÁUSULA SÉTIMA', style="Paragraph2")
    p16.alignment = 3
    
    p17 = Document.add_paragraph(f'\tA sociedade representada poderá ser administrada por sócios ou poderá nomear procurador podendo este ser ou não sócio para representá-la, determinando na procuração, o prazo e a finalidade específica. ', style="Paragraph")
    p17.alignment = 3 
    
    p18 = Document.add_paragraph(f'\nCLÁUSULA OITAVA', style="Paragraph2")
    p18.alignment = 3    
    
    p19 = Document.add_paragraph(f'\tA administração da sociedade caberá ao único (a) sócio (a) administrador (a) {nome_empresario}, acima qualificado (a) sendo exercida isoladamente, com os poderes e atribuições de representação ativa e passiva na sociedade, judicial e extrajudicial, podendo praticar todos os atos compreendidos no objeto social, sempre de interesse da sociedade, autorizado o uso do nome empresarial, vedado, no entanto, fazê-lo em atividades estranhas ao interesse social ou assumir obrigações seja em favor de qualquer dos quotistas ou de terceiros, bem como onerar ou alienar bens imóveis da sociedade, sem autorização do(s) outro(s) sócio(s). ', style="Paragraph")
    p19.alignment = 3 
    
    p20 = Document.add_paragraph(f'\nCLÁUSULA DÉCIMA', style="Paragraph2")
    p20.alignment = 3
        
    p21 = Document.add_paragraph(f'\tNos quatro meses seguintes ao término do exercício social, os sócios deliberarão sobre as contas e designarão administrador(es) quando for o caso. ', style="Paragraph")
    p21.alignment = 3 
       
    p22 = Document.add_paragraph(f'\nCLÁUSULA DÉCIMA PRIMEIRA', style="Paragraph2")
    p22.alignment = 3 
    
    p23 = Document.add_paragraph(f'\tO (a) sócio (a) poderá ser excluído (a), quando a maioria dos sócios, representativa de mais da metade do capital social, entender que um ou mais sócios estão pondo em risco a continuidade da empresa, em virtude de atos de inegável gravidade, mediante alteração do contrato social. A exclusão somente poderá ser determinada em reunião ou assembleia especialmente convocada para esse fim, ciente o acusado em tempo hábil para permitir seu comparecimento e o exercício do direito de defesa. ', style="Paragraph")
    p23.alignment = 3 
              
    p24 = Document.add_paragraph(f'\nCLÁUSULA DÉCIMA SEGUNDA', style="Paragraph2")
    p24.alignment = 3    
    
    p25 = Document.add_paragraph(f'\tO sócio {nome_empresario}, acima qualificado (a) poderá fixar uma retirada mensal, a título de "pro labore", observadas as disposições regulamentares pertinentes. ', style="Paragraph")
    p25.alignment = 3
        
    p26 = Document.add_paragraph(f'\nCLÁUSULA DÉCIMA TERCEIRA', style="Paragraph2")
    p26.alignment = 3    
    
    p27 = Document.add_paragraph(f'\tFalecendo o sócio, a sociedade continuará suas atividades, e os herdeiros e/ou sucessores do sócio falecido poderão ser admitidos como sócios se aprovado pela totalidade dos remanescentes. Não sendo aprovado o ingresso dos herdeiros e/ou sucessores na Sociedade ou, sendo aprovado, caso inexista o interesse destes em se tornarem sócios, o valor de seus haveres será apurado e liquidado com base na situação patrimonial da sociedade, à data da resolução, verificada em balanço especialmente levantado. ', style="Paragraph")
    p27.alignment = 3
    
    p28 = Document.add_paragraph(f'\nCLÁUSULA DÉCIMA QUARTA', style="Paragraph2")
    p28.alignment = 3    
    
    p29 = Document.add_paragraph(f'\tO(s) Administrador(es) declara(m), sob as penas da lei, de que não está(ão) impedidos de exercer a administração da sociedade, por lei especial, ou em virtude de condenação criminal, ou por se encontrar(em) sob os efeitos dela, a pena que vede, ainda que temporariamente, o acesso a cargos públicos; ou por crime falimentar, de prevaricação, peita ou suborno, concussão, peculato, ou contra a economia popular, contra o sistema financeiro nacional, contra normas de defesa da concorrência, contra as relações de consumo, fé pública, ou a propriedade.  ', style="Paragraph")
    p29.alignment = 3
            
    p30 = Document.add_paragraph(f'\nCLÁUSULA DÉCIMA QUINTA', style="Paragraph2")
    p30.alignment = 3  
    
    p31 = Document.add_paragraph(f'\tFica eleito o foro de São Paulo para o exercício e o cumprimento dos direitos e obrigações resultantes deste contrato. ', style="Paragraph")
    p31.alignment = 3
    
    p32 = Document.add_paragraph(f'\n\tE por estarem assim justos e contratados assinam o presente instrumento em 3 vias. ', style="Paragraph")
    p32.alignment = 3
             
    p33 = Document.add_paragraph(f'São Paulo, __ de __________ de 2022', style="Paragraph")
    p33.alignment = 2
    
    p34 = Document.add_paragraph(f'{nome_empresario} ', style="Paragraph2")
    p34.alignment = 3
    
    p35 = Document.add_paragraph(f'Sócio(a) Administrador(a) ', style="Paragraph4")
    p35.alignment = 3
    
    p36 = Document.add_paragraph(f'TESTEMUNHAS', style="Paragraph")
    p36.alignment = 1
    
    p37 = Document.add_paragraph(f'\nEMERSON DE OLIVEIRA MELO				PAULO BORGES DE OLIVEIRA ', style="Paragraph")
    p37.alignment = 3
    
    p38 = Document.add_paragraph(f'RG. 38.724.799-3 SSP/SP					RG. 19.940.163-9 SSP/SP ', style="Paragraph4")
    p38.alignment = 3
                                    
    Document.save(f'{pasta_nova} CONTRATO {nome_novo}.docx')
    os.startfile(f'{pasta_nova} CONTRATO {nome_novo}.docx')

cnpj = 39693485000181 #input('Qual CNPJ: ')
x = dados_da_receita(cnpj)
print(x)
nome_cnpj = x[0]['nome']
endereco_cnpj = x[1]
y = transformacao_ei_em_ltda(cnpj, nome_cnpj, endereco_cnpj)